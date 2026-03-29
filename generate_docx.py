"""Generate implemented_features.docx from implementation summary."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

style_h1 = doc.styles['Heading 1']
style_h1.font.size = Pt(16)
style_h1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

style_h2 = doc.styles['Heading 2']
style_h2.font.size = Pt(13)
style_h2.font.color.rgb = RGBColor(0x2D, 0x4A, 0x7A)

style_h3 = doc.styles['Heading 3']
style_h3.font.size = Pt(11)
style_h3.font.color.rgb = RGBColor(0x3D, 0x5A, 0x8A)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Inches(0.3)


# ═══════════════════════════════════════════
# DOCUMENT CONTENT
# ═══════════════════════════════════════════

doc.add_heading('DeepFocus Implementation Summary', level=0)
p = doc.add_paragraph('Prepared: 2026-03-29. All code implemented, built, and verified.')
p.runs[0].italic = True

# ── 1. System Overview ──
doc.add_heading('1. System Overview', level=1)
doc.add_paragraph(
    'DeepFocus is a multimodal, browser-native concentration detection system that combines '
    'three signal modalities \u2014 visual (webcam), behavioural (keyboard/mouse), and contextual '
    '(browser APIs) \u2014 to estimate user focus in real time. All processing occurs locally in '
    'the browser; no images or raw data leave the device.'
)
doc.add_paragraph(
    'Tech stack: React 19 + Vite (frontend), Django REST Framework (backend), '
    'MediaPipe Face Mesh (vision), XGBoost/LSTM (ML), TensorFlow.js (browser inference).'
)

# ── 2. Architecture ──
doc.add_heading('2. System Architecture', level=1)
add_code_block(
    'Browser (Client-Side)\n'
    '\u251c\u2500 MediaPipe Face Mesh (468 landmarks)  \u2192  8 visual features\n'
    '\u251c\u2500 DOM Event Listeners                   \u2192  7 behavioural features\n'
    '\u251c\u2500 Page Visibility + Focus/Blur APIs     \u2192  5 contextual features\n'
    '\u251c\u2500 Temporal Aggregator (EMA, trend)      \u2192  4 temporal features\n'
    '\u2502\n'
    '\u251c\u2500 Feature Vector (24 continuous features)\n'
    '\u2502         \u2502\n'
    '\u2502         \u251c\u2500 Rule-Based Scorer (fallback)\n'
    '\u2502         \u2514\u2500 TF.js ML Model (<1MB)\n'
    '\u2502              \u2502\n'
    '\u2502              \u2514\u2192 focus_score (0-100)\n'
    '\u2502\n'
    '\u251c\u2500 ESM Self-Report Popup (ground truth)\n'
    '\u2502\n'
    '\u2514\u2500 REST API \u2192 Django Backend (SQLite)\n'
    '               \u251c\u2500 SessionEvent (24 ML features per 2s sample)\n'
    '               \u251c\u2500 SelfReport (ESM + post-session)\n'
    '               \u2514\u2500 ML Export Endpoint (CSV/JSON)'
)

# ── 3. Feature Engineering ──
doc.add_heading('3. Feature Engineering (24 features, 4 modalities)', level=1)

# 3.1 Visual
doc.add_heading('3.1 Visual Features (8) \u2014 MediaPipe Face Mesh', level=2)
doc.add_paragraph(
    'Upgraded from BlazeFace (bounding box only) to FaceLandmarker with 468 3D facial landmarks. '
    'Head pose is extracted from the facial transformation matrix via ZYX Euler decomposition. '
    'Eye Aspect Ratio uses the formula from Soukupov\u00e1 & \u010cech (2016). '
    'Gaze direction is computed from iris centre landmarks (468, 473) normalised within the eye opening.'
)
add_table(
    ['Feature', 'Method', 'Range'],
    [
        ['head_yaw', 'Euler angles from 4\u00d74 facial transformation matrix', '-90\u00b0 to +90\u00b0'],
        ['head_pitch', 'Same matrix decomposition', '-90\u00b0 to +90\u00b0'],
        ['head_roll', 'Same matrix decomposition', '-90\u00b0 to +90\u00b0'],
        ['ear_left', 'EAR = (|p2-p6|+|p3-p5|) / (2\u00b7|p1-p4|), 6 eye landmarks', '0.0 \u2013 0.5'],
        ['ear_right', 'Same formula, right eye landmarks', '0.0 \u2013 0.5'],
        ['gaze_x', 'Iris centre position normalised within eye opening', '-1.0 to +1.0'],
        ['gaze_y', 'Same, vertical axis', '-1.0 to +1.0'],
        ['face_confidence', 'Detection confidence from FaceLandmarker', '0.0 \u2013 1.0'],
    ]
)
doc.add_paragraph(
    'Implementation: useFaceDetection.js \u2014 runs at 10 FPS via requestAnimationFrame with GPU delegation. '
    'Falls back to geometric pose estimation if the transformation matrix is unavailable.'
)

# 3.2 Behavioural
doc.add_heading('3.2 Behavioural Features (7) \u2014 Sliding Window Aggregation', level=2)
doc.add_paragraph(
    'All computed over a 30-second sliding window using circular ring buffers (15 slots \u00d7 2s sampling interval).'
)
add_table(
    ['Feature', 'Unit', 'Description'],
    [
        ['keystroke_rate', 'keys/sec', 'Keydown events per second over 30s window'],
        ['mouse_velocity', 'px/sec', 'Total mouse displacement / window duration'],
        ['mouse_distance', 'px', 'Cumulative mouse travel distance in window'],
        ['click_rate', 'clicks/sec', 'Click events per second'],
        ['scroll_rate', 'events/sec', 'Scroll events per second'],
        ['idle_duration', 'seconds', 'Continuous time since last input event'],
        ['activity_level', '0\u20131', 'Normalised composite of all input types'],
    ]
)
doc.add_paragraph(
    'Implementation: useBehaviourSignals.js \u2014 individual event handlers for keydown, mousemove, click, scroll, '
    'touchstart. Per-sample counters reset every 2 seconds; ring buffers aggregate over the 30s window.'
)

# 3.3 Contextual
doc.add_heading('3.3 Contextual Features (5) \u2014 Browser APIs', level=2)
add_table(
    ['Feature', 'Source API', 'Description'],
    [
        ['tab_switch_count', 'visibilitychange', 'Tab-away events in last 5 minutes'],
        ['window_blur_count', 'window blur', 'Application-level focus losses in last 5 minutes'],
        ['time_since_tab_return', 'visibilitychange / focus', 'Seconds since user last returned to tab'],
        ['session_elapsed_ratio', 'Timer', 'Session progress as fraction (0.0\u20131.0)'],
        ['time_of_day', 'Date', 'Current hour (0\u201323), circadian proxy'],
    ]
)
doc.add_paragraph(
    'Implementation: useContextSignals.js \u2014 event timestamps stored in arrays, pruned to 5-minute window each sample.'
)

# 3.4 Temporal
doc.add_heading('3.4 Temporal Features (4) \u2014 Time-Series Derived', level=2)
add_table(
    ['Feature', 'Window', 'Method'],
    [
        ['focus_ema_30s', '30 seconds', 'Exponential Moving Average (\u03b1 = 0.118)'],
        ['focus_ema_5min', '5 minutes', 'Exponential Moving Average (\u03b1 = 0.013)'],
        ['focus_trend', '5 minutes', 'Linear regression slope over score history'],
        ['distraction_burst_count', '5 minutes', 'Count of consecutive low-score (<50) streaks \u22656 seconds'],
    ]
)
doc.add_paragraph(
    'Implementation: useTemporalFeatures.js \u2014 updates on each new score sample (every 2 seconds). '
    'The linear slope is computed via ordinary least squares over the 5-minute score history buffer.'
)

# ── 4. Scoring System ──
doc.add_heading('4. Scoring System', level=1)

doc.add_heading('4.1 Rule-Based Baseline (fallback)', level=2)
doc.add_paragraph('The original rule-based scorer is retained as a deterministic fallback:')
add_code_block(
    'Camera ON:  score = 100 - 40\u00b7(face_missing) - 30\u00b7(looking_away) - 20\u00b7(tab_hidden) - 10\u00b7(idle)\n'
    'Camera OFF: score = 100 - 67\u00b7(tab_hidden) - 33\u00b7(idle)'
)

doc.add_heading('4.2 ML-Based Scoring', level=2)
doc.add_paragraph(
    'computeFocusScoreML() first checks whether the TF.js model is loaded. If available, it runs inference '
    'on the 24-feature vector and returns the predicted score. Otherwise, it derives boolean signals from '
    'continuous features and delegates to the rule-based scorer.'
)
doc.add_paragraph(
    'assembleFeatureVector() collects all 24 features from the four hooks into a flat object whose keys '
    'match the backend SessionEvent model field names, ensuring consistency between training and inference.'
)

# ── 5. Ground Truth Collection ──
doc.add_heading('5. Ground Truth Collection', level=1)

doc.add_heading('5.1 Experience Sampling Method (ESM)', level=2)
doc.add_paragraph(
    'ESMPopup.jsx \u2014 a small, non-intrusive overlay in the bottom-right corner during active sessions. '
    'It triggers at random intervals between 3 and 8 minutes apart, asking "How focused are you right now?" '
    'on a 1\u20135 Likert scale. The popup auto-dismisses after 10 seconds if the user does not respond. '
    'Responses are stored via POST /api/sessions/<id>/reports/ with report_type "esm".'
)

doc.add_heading('5.2 Post-Session Self-Report', level=2)
doc.add_paragraph(
    'Added to ReportPage.jsx, displayed immediately after a session ends. The user rates their overall focus '
    'on a 1\u201310 scale. Stored with report_type "post_session". Both ESM and post-session labels are stored '
    'in the SelfReport model (session foreign key, timestamp, report_type, score).'
)

# ── 6. Backend Extensions ──
doc.add_heading('6. Backend Extensions', level=1)

doc.add_heading('6.1 Schema Changes', level=2)
doc.add_paragraph(
    'The SessionEvent model was extended with 23 nullable float/integer fields corresponding to all ML features. '
    'All new fields use null=True, blank=True to maintain backward compatibility with sessions recorded without '
    'camera or before the upgrade. A new SelfReport model was added for ground truth storage.'
)
doc.add_paragraph('Migrations: 0003_sessionevent_ml_features, 0004_selfreport.')

doc.add_heading('6.2 New API Endpoints', level=2)
add_table(
    ['Endpoint', 'Method', 'Purpose'],
    [
        ['/api/sessions/<id>/reports/', 'POST, GET', 'Submit and retrieve self-reports for a session'],
        ['/api/ml/export/', 'GET', 'Export all events as CSV (default) or JSON (?format=json) with date filtering'],
    ]
)

# ── 7. ML Training Pipeline ──
doc.add_heading('7. ML Training Pipeline', level=1)
doc.add_paragraph(
    'The ML pipeline is located in the ml/ directory and provides a complete workflow from data export '
    'to browser deployment.'
)

add_table(
    ['Script', 'Purpose'],
    [
        ['export_data.py', 'Extract events + self-reports from SQLite \u2192 CSV files'],
        ['feature_engineering.py', 'Clean data, fill NaN defaults, Z-score normalise, align ESM labels (\u00b15s window), save scaler parameters'],
        ['train_xgboost.py', 'XGBoost regressor (MAE target) + 3-class classifier + ablation study across 6 feature-set combinations'],
        ['train_lstm.py', 'Conv1D + LSTM hybrid on 60-second sliding windows (30 timesteps \u00d7 24 features)'],
        ['evaluate.py', 'Generate thesis figures: feature importance, ablation chart, rule-based vs ML scatter, comparison table'],
        ['export_tfjs.py', 'Knowledge distillation (XGBoost \u2192 small MLP \u2192 TF.js) or direct LSTM conversion'],
    ]
)

doc.add_heading('7.1 XGBoost Configuration', level=2)
doc.add_paragraph(
    '200 estimators, max_depth=6, learning_rate=0.1, subsample=0.8, colsample_bytree=0.8. '
    'Cross-validation uses GroupKFold by session_id to prevent data leakage across sessions. '
    'Outputs: regression metrics (MAE, RMSE), classification metrics (accuracy, F1-macro), '
    'and ranked feature importance.'
)

doc.add_heading('7.2 LSTM Architecture', level=2)
add_code_block(
    'Input(30, 24) \u2192 Conv1D(64, k=3) \u2192 Conv1D(64, k=3) \u2192 MaxPool(2) \u2192 Dropout(0.2)\n'
    '             \u2192 LSTM(64) \u2192 Dropout(0.3) \u2192 Dense(32, ReLU) \u2192 Dense(1)'
)
doc.add_paragraph(
    'Trained with EarlyStopping (patience=10) and ReduceLROnPlateau. Train/validation split by session (80/20).'
)

doc.add_heading('7.3 Ablation Study Design', level=2)
doc.add_paragraph('Six feature-set combinations are tested to quantify each modality\u2019s contribution:')
add_table(
    ['#', 'Feature Set', 'Number of Features'],
    [
        ['1', 'Visual only', '8'],
        ['2', 'Behavioural only', '7'],
        ['3', 'Contextual only', '5'],
        ['4', 'Temporal only', '4'],
        ['5', 'Visual + Behavioural', '15'],
        ['6', 'All features (multimodal)', '24'],
    ]
)

# ── 8. Browser-Side ML Inference ──
doc.add_heading('8. Browser-Side ML Inference', level=1)
doc.add_paragraph(
    'FocusModel.js loads TensorFlow.js from CDN at runtime (not bundled, avoiding a 3MB dependency). '
    'The loading sequence is: (1) fetch model metadata (model_meta.json) containing the feature list, '
    '(2) inject TF.js script tag from CDN, (3) load the TF.js graph model (model.json + weight shards), '
    '(4) on each prediction, apply Z-score normalisation using saved scaler parameters and return a '
    'clamped score (0\u2013100).'
)
doc.add_paragraph(
    'Graceful fallback: if any step fails (model not deployed, CDN unreachable, prediction error), '
    'the system silently reverts to rule-based scoring. The target model size is under 1 MB '
    '(knowledge-distilled MLP: Input \u2192 32 \u2192 16 \u2192 1).'
)

# ── 9. Data Flow ──
doc.add_heading('9. Data Flow Summary', level=1)
add_code_block(
    'Session Running (every 2 seconds):\n'
    '  Face Mesh \u2192 {headYaw, headPitch, headRoll, earL, earR, gazeX, gazeY, confidence}\n'
    '  DOM Events \u2192 {keystrokeRate, mouseVelocity, mouseDist, clickRate, scrollRate, idleDuration, activityLevel}\n'
    '  Browser APIs \u2192 {tabSwitchCount, windowBlurCount, timeSinceTabReturn, sessionElapsedRatio, timeOfDay}\n'
    '  Temporal \u2192 {focusEma30s, focusEma5min, focusTrend, distractionBurstCount}\n'
    '       \u2193\n'
    '  assembleFeatureVector() \u2192 24 named features\n'
    '       \u2193\n'
    '  computeFocusScore() [rule-based] or computeFocusScoreML() [TF.js]\n'
    '       \u2193\n'
    '  Event stored: {timestamp, focus_score, 4 legacy booleans, 24 ML features}\n'
    '       \u2193\n'
    '  On session end: bulk upload to backend \u2192 stored in SessionEvent table'
)

# ── 10. Files Changed ──
doc.add_heading('10. Files Changed and Created', level=1)

doc.add_heading('Frontend (10 files)', level=2)
add_table(
    ['File', 'Action', 'Description'],
    [
        ['useFaceDetection.js', 'Rewritten', 'BlazeFace \u2192 Face Mesh 468 landmarks, PnP head pose, EAR, iris gaze'],
        ['useBehaviourSignals.js', 'Rewritten', '30s sliding-window ring buffers, keystroke/mouse/click/scroll rates'],
        ['useContextSignals.js', 'New', 'Tab switch count, window blur count, session progress'],
        ['useTemporalFeatures.js', 'New', 'EMA 30s/5min, focus trend (linear regression), distraction bursts'],
        ['useSession.js', 'Rewritten', 'Wires all hooks, sends ML features in events'],
        ['scoring.js', 'Extended', 'assembleFeatureVector() + computeFocusScoreML()'],
        ['FocusModel.js', 'New', 'CDN-loaded TF.js inference with graceful fallback'],
        ['ESMPopup.jsx', 'New', 'Random self-report popup (3\u20138 min interval, 1\u20135 scale)'],
        ['SessionPage.jsx', 'Updated', 'Face Mesh info display, ESM popup integration'],
        ['ReportPage.jsx', 'Updated', 'Post-session 1\u201310 self-report UI'],
    ]
)

doc.add_heading('Backend (4 files)', level=2)
add_table(
    ['File', 'Action', 'Description'],
    [
        ['models.py', 'Extended', '+23 ML feature fields on SessionEvent, +SelfReport model'],
        ['serializers.py', 'Extended', 'ML features in SessionEventSerializer, +SelfReportSerializer'],
        ['views.py', 'Extended', '+SessionReportsView, +MLExportView (CSV/JSON)'],
        ['urls.py', 'Extended', '+/sessions/<id>/reports/, +/ml/export/'],
    ]
)

doc.add_heading('ML Pipeline (6 new files)', level=2)
add_table(
    ['File', 'Purpose'],
    [
        ['ml/export_data.py', 'SQLite/API \u2192 CSV export'],
        ['ml/feature_engineering.py', 'Clean, normalise, align ESM labels'],
        ['ml/train_xgboost.py', 'XGBoost regression + 3-class + ablation'],
        ['ml/train_lstm.py', 'Conv1D + LSTM temporal model'],
        ['ml/evaluate.py', 'Thesis figures: feature importance, ablation, comparison'],
        ['ml/export_tfjs.py', 'Knowledge distillation \u2192 TF.js conversion'],
    ]
)

# ── 11. Remaining Work ──
doc.add_heading('11. Remaining Work (Step 6: Evaluation)', level=1)
doc.add_paragraph('The evaluation phase requires real user data from testing sessions:')
p = doc.add_paragraph()
items = [
    'Deploy application and conduct user testing (target: 1,000+ ESM-labelled events from 20+ users)',
    'Run ML pipeline: export_data.py \u2192 feature_engineering.py \u2192 train_xgboost.py \u2192 evaluate.py',
    'Export trained model to browser: export_tfjs.py',
    'Generate thesis figures from ml/results/',
    'Conduct user study: SUS (System Usability Scale) questionnaire + semi-structured interviews',
    'Compare perceived focus (self-report) vs ML-predicted focus',
]
for item in items:
    doc.add_paragraph(item, style='List Number')

# ── Save ──
output_path = '/Users/juhyeongpark/Desktop/Individual-Project/DeepFocus/implemented_features.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
